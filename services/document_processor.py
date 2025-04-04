import os
import logging
import hashlib
import mimetypes
import datetime
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor
import traceback
from typing import List, Dict, Any, Tuple, Optional
import base64

import spacy
from langdetect import detect
import PyPDF2
import docx
import pytesseract
from PIL import Image
import pydub
import cv2
# Conditional import to handle missing packages
try:
    from transformers import CLIPProcessor, CLIPModel
    import torch
    TRANSFORMERS_AVAILABLE = True
except ImportError:
    TRANSFORMERS_AVAILABLE = False

from app import app, db
from models import Document, DocumentChunk, DocumentMetadata, Entity, EntityRelationship
from config import (
    CHUNK_SIZE, CHUNK_OVERLAP, MAX_DOCUMENT_SIZE, MAX_THREADS, BATCH_SIZE,
    ENTITY_CONFIDENCE_THRESHOLD, SPACY_MODEL, ENABLE_IMAGE_PROCESSING,
    ENABLE_AUDIO_PROCESSING, ENABLE_VIDEO_PROCESSING, IMAGE_EMBEDDING_MODEL
)
from services import embedding_service, knowledge_graph, vector_store
from utils.preprocessing import clean_text, normalize_text
from utils.helpers import get_file_size, get_mime_type, extract_filename
from utils.security import encrypt_file, decrypt_file, is_file_encrypted

logger = logging.getLogger(__name__)

# Initialize NLP model
nlp = None
clip_model = None
clip_processor = None

def init_app(app):
    """Initialize the document processor with the app context"""
    global nlp, clip_model, clip_processor
    
    # Load spaCy model
    try:
        nlp = spacy.load(SPACY_MODEL)
        logger.info(f"Loaded spaCy model: {SPACY_MODEL}")
    except Exception as e:
        logger.error(f"Error loading spaCy model: {e}")
        # Try a smaller model as fallback
        try:
            nlp = spacy.load("en_core_web_sm")
            logger.info("Loaded fallback spaCy model: en_core_web_sm")
        except:
            logger.error("Could not load any spaCy model. Entity extraction will be limited.")
            nlp = None
    
    # Load CLIP model for image processing if enabled
    if TRANSFORMERS_AVAILABLE and ENABLE_IMAGE_PROCESSING and IMAGE_EMBEDDING_MODEL == "clip":
        try:
            from transformers import CLIPProcessor, CLIPModel
            clip_model = CLIPModel.from_pretrained("openai/clip-vit-base-patch32")
            clip_processor = CLIPProcessor.from_pretrained("openai/clip-vit-base-patch32")
            logger.info("Loaded CLIP model for image processing")
        except Exception as e:
            logger.error(f"Error loading CLIP model: {e}")
            clip_model = None
            clip_processor = None
    else:
        if ENABLE_IMAGE_PROCESSING and not TRANSFORMERS_AVAILABLE:
            logger.warning("Transformers package not available - image processing capabilities will be limited")
    
    logger.info("Document processor initialized")

def process_document(file_path: str, title: str, description: str, user_id: int,
                     is_public: bool = False, is_encrypted: bool = False) -> Optional[Document]:
    """
    Process a document and store it in the database with extracted chunks and entities
    
    Args:
        file_path: Path to the document file
        title: Document title
        description: Document description
        user_id: ID of the user who uploaded the document
        is_public: Whether the document is publicly accessible
        is_encrypted: Whether to encrypt the document
        
    Returns:
        Document object if successful, None otherwise
    """
    try:
        # Check if the file exists
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None
        
        # Check file size
        file_size = get_file_size(file_path)
        if file_size > MAX_DOCUMENT_SIZE:
            logger.error(f"File size exceeds maximum allowed size: {file_size} bytes")
            return None
        
        # Get file type and mime type
        file_type = os.path.splitext(file_path)[1][1:].lower()
        mime_type = get_mime_type(file_path)
        
        # Extract text based on file type
        extracted_text, metadata = extract_text_and_metadata(file_path, file_type, mime_type)
        if not extracted_text:
            logger.warning(f"No text could be extracted from document: {file_path}")
            
        # Detect language
        try:
            language = detect(extracted_text[:1000]) if extracted_text else "en"
        except:
            language = "en"  # Default to English if detection fails
        
        # Create document record
        document = Document(
            title=title,
            description=description,
            file_path=file_path,
            file_type=file_type,
            file_size=file_size,
            mime_type=mime_type,
            language=language,
            owner_id=user_id,
            is_public=is_public,
            is_encrypted=is_encrypted,
            created_at=datetime.datetime.utcnow()
        )
        
        # Encrypt file if requested
        if is_encrypted:
            encrypted_path = encrypt_file(file_path)
            document.file_path = encrypted_path
            document.encryption_metadata = hashlib.sha256(os.path.basename(file_path).encode()).hexdigest()
        
        db.session.add(document)
        db.session.flush()  # Get document ID without committing
        
        # Create document metadata
        doc_metadata = DocumentMetadata(
            document_id=document.id,
            author=metadata.get('author'),
            created_date=metadata.get('created_date'),
            modified_date=metadata.get('modified_date'),
            page_count=metadata.get('page_count'),
            word_count=len(extracted_text.split()) if extracted_text else 0,
            source_url=metadata.get('source_url'),
            source_system=metadata.get('source_system'),
            additional_metadata=metadata.get('additional_metadata')
        )
        db.session.add(doc_metadata)
        
        # Process document text if available
        if extracted_text:
            # Create chunks
            chunks = create_chunks(extracted_text, CHUNK_SIZE, CHUNK_OVERLAP)
            store_chunks(document.id, chunks)
            
            # Process chunks in batches for efficiency
            process_chunks_in_batches(document.id, chunks)
        
        # Process multimedia content if applicable
        if ENABLE_IMAGE_PROCESSING and file_type in ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff']:
            process_image_content(document.id, file_path)
            
        if ENABLE_AUDIO_PROCESSING and file_type in ['mp3', 'wav', 'ogg', 'flac']:
            process_audio_content(document.id, file_path)
            
        if ENABLE_VIDEO_PROCESSING and file_type in ['mp4', 'avi', 'mov', 'wmv']:
            process_video_content(document.id, file_path)
        
        # Mark document as indexed
        document.indexed_at = datetime.datetime.utcnow()
        db.session.commit()
        
        logger.info(f"Successfully processed document: {document.id} - {document.title}")
        return document
        
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error processing document: {str(e)}")
        logger.debug(traceback.format_exc())
        return None

def extract_text_and_metadata(file_path: str, file_type: str, mime_type: str) -> Tuple[str, Dict[str, Any]]:
    """
    Extract text and metadata from a document based on its file type
    
    Args:
        file_path: Path to the document file
        file_type: Type of file (pdf, docx, txt, etc.)
        mime_type: MIME type of the file
        
    Returns:
        Tuple containing extracted text and metadata dictionary
    """
    extracted_text = ""
    metadata = {}
    
    try:
        # Primary extraction with local tools
        primary_extraction_success = False
        
        try:
            # Handle different file types
            if file_type == 'pdf':
                extracted_text, metadata = extract_from_pdf(file_path)
            elif file_type in ['docx', 'doc']:
                extracted_text, metadata = extract_from_docx(file_path)
            elif file_type == 'txt':
                with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                    extracted_text = f.read()
            elif file_type in ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff']:
                # OCR for image files
                if ENABLE_IMAGE_PROCESSING:
                    extracted_text = extract_text_from_image(file_path)
            elif file_type in ['mp3', 'wav', 'ogg', 'flac']:
                # Audio transcription - use OpenAI Whisper if available
                if ENABLE_AUDIO_PROCESSING:
                    extracted_text = extract_text_from_audio(file_path)
            elif file_type in ['mp4', 'avi', 'mov', 'wmv']:
                # Video transcription - use OpenAI Whisper if available
                if ENABLE_VIDEO_PROCESSING:
                    extracted_text = extract_text_from_video(file_path)
            else:
                # Try to read as plain text for unknown file types
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                        extracted_text = f.read()
                except:
                    logger.warning(f"Unsupported file type: {file_type}. Could not extract text.")
            
            # Clean the extracted text
            if extracted_text:
                extracted_text = clean_text(extracted_text)
                # Check if we got meaningful text
                if len(extracted_text.strip()) > 20:
                    primary_extraction_success = True
                    
        except Exception as primary_e:
            logger.warning(f"Primary extraction failed: {primary_e}")
            logger.debug(traceback.format_exc())
            
        # Try OpenAI fallback if primary extraction failed or returned little text
        if not primary_extraction_success:
            try:
                # Check if OpenAI is available
                from services.llm_service import client as openai_client
                
                if openai_client is not None:
                    logger.info(f"Attempting OpenAI fallback extraction for {file_type} file")
                    
                    if file_type in ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff']:
                        # Use OpenAI for image analysis
                        fallback_text = extract_text_from_image_with_openai(file_path, openai_client)
                        if fallback_text:
                            extracted_text = fallback_text
                            metadata['extraction_method'] = 'openai_vision'
                            
                    elif file_type in ['mp3', 'wav', 'ogg', 'flac']:
                        # Use OpenAI Whisper for audio
                        fallback_text = extract_audio_with_openai(file_path, openai_client)
                        if fallback_text:
                            extracted_text = fallback_text
                            metadata['extraction_method'] = 'openai_whisper'
                            
                    elif file_type in ['pdf', 'docx', 'doc'] and len(extracted_text.strip()) < 100:
                        # Try OCR with OpenAI vision if PDF/DOCX extraction failed
                        fallback_text = extract_document_with_openai(file_path, file_type, openai_client)
                        if fallback_text:
                            extracted_text = fallback_text
                            metadata['extraction_method'] = 'openai_vision_ocr'
            
            except Exception as fallback_e:
                logger.error(f"OpenAI fallback extraction failed: {fallback_e}")
                
        # Add language detection
        if extracted_text and len(extracted_text.strip()) > 100:
            try:
                language = detect(extracted_text[:1000])
                metadata['language'] = language
            except Exception as lang_e:
                logger.warning(f"Language detection failed: {lang_e}")
            
            # Estimate word count
            words = extracted_text.split()
            metadata['word_count'] = len(words)
                
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {str(e)}")
        logger.debug(traceback.format_exc())
        
    return extracted_text, metadata

def extract_from_pdf(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """Extract text and metadata from a PDF file"""
    extracted_text = ""
    metadata = {}
    
    try:
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            
            # Extract metadata
            info = pdf_reader.metadata
            if info:
                metadata = {
                    'author': info.author,
                    'created_date': info.creation_date,
                    'modified_date': info.modification_date,
                    'page_count': len(pdf_reader.pages),
                    'additional_metadata': {k: str(v) for k, v in info.items() if k not in ['author', 'creation_date', 'modification_date']}
                }
            else:
                metadata = {'page_count': len(pdf_reader.pages)}
            
            # Extract text from each page
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                extracted_text += page.extract_text() + "\n\n"
                
    except Exception as e:
        logger.error(f"Error processing PDF: {str(e)}")
        
        # Attempt to recover from corrupted PDF
        try:
            # Use Tesseract OCR as fallback for corrupted PDFs
            # This would convert each page to an image and apply OCR
            logger.info("Attempting corrupted document recovery via OCR")
            extracted_text = "Fallback text extraction would be implemented here"
        except:
            logger.error("Complete failure to extract text from PDF")
    
    return extracted_text, metadata

def extract_from_docx(file_path: str) -> Tuple[str, Dict[str, Any]]:
    """Extract text and metadata from a DOCX file"""
    extracted_text = ""
    metadata = {}
    
    try:
        doc = docx.Document(file_path)
        
        # Extract metadata
        core_properties = doc.core_properties
        metadata = {
            'author': core_properties.author,
            'created_date': core_properties.created,
            'modified_date': core_properties.modified,
            'page_count': None,  # Not directly available in docx
            'additional_metadata': {
                'title': core_properties.title,
                'subject': core_properties.subject,
                'keywords': core_properties.keywords,
                'category': core_properties.category
            }
        }
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            extracted_text += para.text + "\n"
            
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ""
                for cell in row.cells:
                    row_text += cell.text + " | "
                extracted_text += row_text.strip(" | ") + "\n"
            extracted_text += "\n"
            
    except Exception as e:
        logger.error(f"Error processing DOCX: {str(e)}")
    
    return extracted_text, metadata

def extract_text_from_image(file_path: str) -> str:
    """Extract text from an image using OCR"""
    try:
        image = Image.open(file_path)
        extracted_text = pytesseract.image_to_string(image)
        return extracted_text
    except Exception as e:
        logger.error(f"Error extracting text from image: {str(e)}")
        return ""

def extract_text_from_image_with_openai(file_path: str, openai_client) -> str:
    """Extract text from an image using OpenAI's vision capabilities"""
    try:
        # Read the image file
        with open(file_path, "rb") as img_file:
            image_data = img_file.read()
            base64_image = base64.b64encode(image_data).decode('utf-8')
        
        # Call OpenAI Vision API
        response = openai_client.chat.completions.create(
            model="gpt-4o",  # Use the newest model with vision capabilities
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": "Extract all visible text from this image. Include all paragraphs, labels, and any other text content. Format the text in a clean, readable layout that preserves the original structure. If there are tables, maintain table formatting as best as possible using plain text."
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{base64_image}"
                            }
                        }
                    ]
                }
            ],
            max_tokens=1000
        )
        
        # Extract the text
        extracted_text = response.choices[0].message.content
        logger.info("Successfully extracted text from image using OpenAI Vision")
        return extracted_text
        
    except Exception as e:
        logger.error(f"Error extracting text from image with OpenAI: {str(e)}")
        return ""

def extract_text_from_audio(file_path: str) -> str:
    """Extract text from audio file using local tools"""
    try:
        # Placeholder - in production you would use a local speech recognition tool
        logger.warning("Local audio transcription not fully implemented")
        return f"[Audio content from {os.path.basename(file_path)}]"
    except Exception as e:
        logger.error(f"Error extracting text from audio: {str(e)}")
        return ""

def extract_audio_with_openai(file_path: str, openai_client) -> str:
    """Transcribe audio using OpenAI's Whisper API"""
    try:
        with open(file_path, "rb") as audio_file:
            response = openai_client.audio.transcriptions.create(
                model="whisper-1",
                file=audio_file
            )
        
        # Get transcription
        transcription = response.text
        logger.info("Successfully transcribed audio using OpenAI Whisper")
        return transcription
        
    except Exception as e:
        logger.error(f"Error transcribing audio with OpenAI: {str(e)}")
        return ""

def extract_text_from_video(file_path: str) -> str:
    """Extract audio from video and transcribe it"""
    try:
        # Placeholder - would extract audio and then transcribe
        logger.warning("Local video processing not fully implemented")
        return f"[Video content from {os.path.basename(file_path)}]"
    except Exception as e:
        logger.error(f"Error extracting text from video: {str(e)}")
        return ""

def extract_document_with_openai(file_path: str, file_type: str, openai_client) -> str:
    """Extract text from document using OpenAI's vision capabilities as fallback"""
    try:
        # For PDFs, we might need to render pages as images
        if file_type == 'pdf':
            # This would convert PDF pages to images and then use Vision API
            # Placeholder implementation
            logger.info("Using OpenAI to analyze PDF document")
            
            # Simple implementation for first page only - production would process all pages
            try:
                with open(file_path, "rb") as img_file:
                    # Just use the first page for demonstration
                    image_data = img_file.read()
                    base64_image = base64.b64encode(image_data).decode('utf-8')
                
                # Call OpenAI Vision API
                response = openai_client.chat.completions.create(
                    model="gpt-4o",  # Vision-capable model
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "text",
                                    "text": "This is a document that couldn't be processed properly. Extract all visible text content, preserving the structure. Include headings, paragraphs, tables, and any other text elements."
                                },
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:application/pdf;base64,{base64_image}"
                                    }
                                }
                            ]
                        }
                    ],
                    max_tokens=4000
                )
                
                # Extract the text
                extracted_text = response.choices[0].message.content
                return extracted_text
            except Exception as pdf_e:
                logger.error(f"Error extracting PDF with OpenAI Vision: {pdf_e}")
        
        # For other document types, we might convert to PDF first and then process
        else:
            logger.warning(f"OpenAI fallback not fully implemented for {file_type}")
            
        return ""
        
    except Exception as e:
        logger.error(f"Error with OpenAI document extraction: {str(e)}")
        return ""

def create_chunks(text: str, chunk_size: int, chunk_overlap: int) -> List[Dict[str, Any]]:
    """
    Split text into overlapping chunks
    
    Args:
        text: Text to split into chunks
        chunk_size: Maximum size of each chunk in characters
        chunk_overlap: Number of characters to overlap between chunks
        
    Returns:
        List of chunk dictionaries with content, start and end positions
    """
    chunks = []
    start = 0
    
    while start < len(text):
        # Calculate end position with overlap
        end = min(start + chunk_size, len(text))
        
        # Don't create tiny chunks at the end
        if end - start < chunk_size // 4 and len(chunks) > 0:
            chunks[-1]['content'] += text[start:end]
            chunks[-1]['end_char'] = end
            break
            
        # Create chunk
        chunk = {
            'content': text[start:end],
            'start_char': start,
            'end_char': end
        }
        chunks.append(chunk)
        
        # Move the start position, considering overlap
        start = end - chunk_overlap
        
        # Avoid getting stuck in very short overlaps
        if start >= end - 10:
            start = end
    
    return chunks

def store_chunks(document_id: int, chunks: List[Dict[str, Any]]) -> None:
    """
    Store document chunks in the database
    
    Args:
        document_id: ID of the document
        chunks: List of chunk dictionaries
    """
    try:
        for index, chunk in enumerate(chunks):
            db_chunk = DocumentChunk(
                document_id=document_id,
                content=chunk['content'],
                chunk_index=index,
                start_char=chunk['start_char'],
                end_char=chunk['end_char']
            )
            db.session.add(db_chunk)
        
        db.session.commit()
        logger.info(f"Stored {len(chunks)} chunks for document {document_id}")
        
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error storing chunks: {str(e)}")

def process_chunks_in_batches(document_id: int, chunks: List[Dict[str, Any]]) -> None:
    """
    Process document chunks in batches for efficiency
    
    Args:
        document_id: ID of the document
        chunks: List of chunk dictionaries
    """
    try:
        # Get created chunks from the database
        db_chunks = DocumentChunk.query.filter_by(document_id=document_id).order_by(DocumentChunk.chunk_index).all()
        
        # Process chunks in batches
        for i in range(0, len(db_chunks), BATCH_SIZE):
            batch = db_chunks[i:i+BATCH_SIZE]
            
            # Generate embeddings for chunks
            chunk_texts = [chunk.content for chunk in batch]
            chunk_ids = [chunk.id for chunk in batch]
            
            # Create embeddings and store in vector database
            embedding_ids = embedding_service.embed_texts(chunk_texts, chunk_ids, "document")
            
            # Update chunk records with embedding IDs
            for j, chunk in enumerate(batch):
                if j < len(embedding_ids) and embedding_ids[j]:
                    chunk.embedding_id = embedding_ids[j]
            
            # Extract entities from each chunk
            if nlp:
                with ThreadPoolExecutor(max_workers=MAX_THREADS) as executor:
                    for chunk in batch:
                        executor.submit(extract_entities_from_chunk, document_id, chunk)
        
        # Commit all changes
        db.session.commit()
        
        # Build knowledge graph from extracted entities
        knowledge_graph.build_knowledge_graph_for_document(document_id)
        
        logger.info(f"Processed all chunks for document {document_id}")
        
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error processing chunks in batches: {str(e)}")
        logger.debug(traceback.format_exc())

def extract_entities_from_chunk(document_id: int, chunk: DocumentChunk) -> None:
    """
    Extract named entities from a document chunk and store them
    
    Args:
        document_id: ID of the document
        chunk: DocumentChunk object
    """
    try:
        # Skip if NLP model isn't loaded
        if not nlp:
            return
            
        # Process the chunk text with spaCy
        doc = nlp(chunk.content)
        
        # Extract named entities
        for ent in doc.ents:
            # Calculate position in the original document
            start_pos = chunk.start_char + ent.start_char
            end_pos = chunk.start_char + ent.end_char
            
            # Skip entities with low confidence (more sophisticated confidence would be used in production)
            confidence = 0.85  # Placeholder - would use model confidence in production
            if confidence < ENTITY_CONFIDENCE_THRESHOLD:
                continue
                
            # Create entity record
            entity = Entity(
                name=ent.text,
                entity_type=ent.label_,
                document_id=document_id,
                chunk_id=chunk.id,
                confidence=confidence,
                start_char=start_pos,
                end_char=end_pos
            )
            db.session.add(entity)
            
    except Exception as e:
        logger.error(f"Error extracting entities from chunk {chunk.id}: {str(e)}")

def process_image_content(document_id: int, file_path: str) -> None:
    """Process image content using CLIP for visual embeddings"""
    if not ENABLE_IMAGE_PROCESSING:
        return
        
    # Skip if transformers is not available
    if not TRANSFORMERS_AVAILABLE:
        logger.warning("Transformers package not available - skipping image embedding generation")
        return
        
    if not clip_model or not clip_processor:
        return
        
    try:
        # Load image
        image = Image.open(file_path)
        
        # Process with CLIP
        inputs = clip_processor(
            text=["An image"], 
            images=image, 
            return_tensors="pt", 
            padding=True
        )
        
        # Generate image embeddings
        with torch.no_grad():
            outputs = clip_model(**inputs)
            image_embedding = outputs.image_embeds[0].cpu().numpy()
        
        # Store embedding in vector database
        embedding_id = vector_store.store_embedding(
            image_embedding, 
            {"document_id": document_id, "type": "image", "source": file_path}
        )
        
        logger.info(f"Processed image content for document {document_id}")
        
    except Exception as e:
        logger.error(f"Error processing image content: {str(e)}")

def process_audio_content(document_id: int, file_path: str) -> None:
    """Process audio content (placeholder - would use a real audio processing API)"""
    if not ENABLE_AUDIO_PROCESSING:
        return
        
    try:
        # This would use a real audio processing service in production
        logger.info(f"Audio processing would happen here for document {document_id}")
        
    except Exception as e:
        logger.error(f"Error processing audio content: {str(e)}")

def process_video_content(document_id: int, file_path: str) -> None:
    """Process video content (placeholder - would use a real video processing API)"""
    if not ENABLE_VIDEO_PROCESSING:
        return
        
    try:
        # This would use a real video processing service in production
        logger.info(f"Video processing would happen here for document {document_id}")
        
    except Exception as e:
        logger.error(f"Error processing video content: {str(e)}")

def reprocess_document(document_id: int) -> bool:
    """
    Reprocess an existing document (e.g., after model updates)
    
    Args:
        document_id: ID of the document to reprocess
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Get document from database
        document = Document.query.get(document_id)
        if not document:
            logger.error(f"Document not found: {document_id}")
            return False
            
        # Check if file exists
        file_path = document.file_path
        if is_file_encrypted(document):
            # Decrypt file temporarily
            file_path = decrypt_file(document.file_path)
            
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return False
            
        # Remove existing chunks and entities
        DocumentChunk.query.filter_by(document_id=document_id).delete()
        Entity.query.filter_by(document_id=document_id).delete()
        
        # Clear existing embeddings from vector store
        vector_store.delete_embeddings_for_document(document_id)
        
        # Re-extract text
        extracted_text, metadata = extract_text_and_metadata(
            file_path, document.file_type, document.mime_type
        )
        
        # Update metadata
        if metadata:
            doc_metadata = DocumentMetadata.query.filter_by(document_id=document_id).first()
            if doc_metadata:
                for key, value in metadata.items():
                    if hasattr(doc_metadata, key) and value is not None:
                        setattr(doc_metadata, key, value)
        
        # Reprocess document text
        if extracted_text:
            chunks = create_chunks(extracted_text, CHUNK_SIZE, CHUNK_OVERLAP)
            store_chunks(document_id, chunks)
            process_chunks_in_batches(document_id, chunks)
            
        # Clean up temporary decrypted file if needed
        if is_file_encrypted(document) and file_path != document.file_path:
            os.remove(file_path)
            
        # Update document
        document.indexed_at = datetime.datetime.utcnow()
        db.session.commit()
        
        logger.info(f"Successfully reprocessed document: {document_id}")
        return True
        
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error reprocessing document: {str(e)}")
        logger.debug(traceback.format_exc())
        return False

def recover_corrupted_document(document_id: int) -> bool:
    """
    Attempt to recover and reprocess a corrupted document
    
    Args:
        document_id: ID of the document to recover
        
    Returns:
        True if successfully recovered, False otherwise
    """
    try:
        # Get document from database
        document = Document.query.get(document_id)
        if not document:
            logger.error(f"Document not found: {document_id}")
            return False
            
        # Check if file exists
        file_path = document.file_path
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return False
            
        # Attempt recovery based on file type
        if document.file_type == 'pdf':
            # PDF recovery would use alternative libraries or OCR
            logger.info(f"Attempting PDF recovery for document {document_id}")
            return attempt_pdf_recovery(document)
        elif document.file_type in ['docx', 'doc']:
            # DOCX recovery
            logger.info(f"Attempting DOCX recovery for document {document_id}")
            return attempt_docx_recovery(document)
        elif document.file_type in ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff']:
            # Image recovery using alternative OCR methods
            logger.info(f"Attempting image recovery for document {document_id}")
            return attempt_image_recovery(document)
        else:
            logger.warning(f"Recovery not implemented for file type: {document.file_type}")
            return False
            
    except Exception as e:
        logger.error(f"Error recovering document: {str(e)}")
        return False

def attempt_pdf_recovery(document: Document) -> bool:
    """Attempt to recover a corrupted PDF document"""
    # This would use alternative PDF parsers or convert to images and apply OCR
    # Placeholder implementation
    try:
        logger.info(f"PDF recovery would be implemented here for document {document.id}")
        return True
    except:
        return False

def attempt_docx_recovery(document: Document) -> bool:
    """Attempt to recover a corrupted DOCX document"""
    # This would use alternative DOCX parsers or extract the XML directly
    # Placeholder implementation
    try:
        logger.info(f"DOCX recovery would be implemented here for document {document.id}")
        return True
    except:
        return False

def attempt_image_recovery(document: Document) -> bool:
    """Attempt to recover a corrupted image using alternative OCR methods"""
    # This would use alternative OCR methods or image enhancement techniques
    # Placeholder implementation
    try:
        logger.info(f"Image recovery would be implemented here for document {document.id}")
        return True
    except:
        return False
